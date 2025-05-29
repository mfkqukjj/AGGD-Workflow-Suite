import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
import threading

class BatchGenerator(tk.Toplevel):
    def __init__(self,master):
        super().__init__(master)
        self.transient(master)  # 标记为主窗口的附属窗口
        self.grab_set()         # 独占焦点
        self.attributes('-topmost', True)  # 强制置顶
        self.title("AGGD-塘厦分公司-客户一人一档批量生成器 by.281303")
        self.geometry('950x700')

        # 初始化变量
        self.template_path = ""
        self.roster_data = []
        self.replace_rules = []
        self.output_dir = ""
        self.auto_dir = tk.BooleanVar(value=True)
        self.column_names = []  # 新增：存储Excel列名

        # GUI组件初始化
        self.create_widgets()
        self._init_replace_rules(initial_rows=5)
        
        self.mainloop()

    def create_widgets(self):
        """创建界面组件"""
        # 文件导入区域
        ttk.Button(self, text="导入模板文件(Word/Excel)", 
                  command=self.load_template).grid(row=0, column=0, padx=5, pady=5, sticky='ew')
        ttk.Button(self, text="导入人员台账(Excel)", 
                  command=self.load_roster).grid(row=0, column=1, padx=5, pady=5, sticky='ew')

        # 替换规则区域
        self.rule_frame = ttk.LabelFrame(self, text="替换规则配置")
        self.rule_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky='nsew')

        # 动态规则管理按钮
        ttk.Button(self.rule_frame, text="添加规则", command=self._add_rule_row).grid(row=0, column=2, padx=5)
        ttk.Button(self.rule_frame, text="删除规则", command=self._remove_rule_row).grid(row=0, column=3, padx=5)

        # 文件名格式
        ttk.Label(self, text="文件名格式:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.filename_pattern = ttk.Entry(self, width=50)
        self.filename_pattern.insert(0, "一人一档（{姓名}{身份证}）")
        self.filename_pattern.grid(row=2, column=1, padx=5, pady=5, sticky='w')

        # 输出目录
        ttk.Checkbutton(self, text="自动创建子目录", 
                       variable=self.auto_dir).grid(row=3, column=0, padx=5, pady=5)
        ttk.Button(self, text="选择输出目录", 
                  command=self.select_output_dir).grid(row=3, column=1, padx=5, pady=5)

        # 进度条
        self.progress = ttk.Progressbar(self, orient='horizontal', 
                                       length=300, mode='determinate')
        self.progress.grid(row=4, column=0, columnspan=2, pady=10)

        # 运行按钮
        ttk.Button(self, text="开始生成", 
                  command=self.start_generation).grid(row=5, column=1, pady=10, sticky='e')
        
    def _init_replace_rules(self, initial_rows=5):
        """初始化替换规则输入区域"""
        ttk.Label(self.rule_frame, text="模板占位符").grid(row=0, column=0, padx=5)
        ttk.Label(self.rule_frame, text="对应台账字段").grid(row=0, column=1, padx=5)
        for i in range(initial_rows):
            self._add_rule_row(row=i+1)


    def _add_rule_row(self, row=None):
        """添加单行规则输入（改为下拉框）"""
        row = len(self.replace_rules)+1 if row is None else row
        placeholder = ttk.Entry(self.rule_frame, width=22)
        field = ttk.Combobox(self.rule_frame, width=20)  # 改为下拉框
        field['values'] = self.column_names  # 设置初始选项
        
        placeholder.grid(row=row, column=0, padx=5, pady=2)
        field.grid(row=row, column=1, padx=5, pady=2)
        self.replace_rules.append((placeholder, field))

    def _remove_rule_row(self):
        """删除最后一条规则"""
        if len(self.replace_rules) > 1:
            for widget in self.replace_rules.pop():
                widget.destroy()

    # ------------------- 文件操作函数 -------------------
    def load_template(self):
        """加载模板文件（支持Excel）"""
        path = filedialog.askopenfilename(
            parent=self,  
            filetypes=[("Word文件", "*.docx"), ("Excel文件", "*.xlsx")])
        if path:
            self.template_path = path
            self.template_type = "docx" if path.endswith(".docx") else "xlsx"
            messagebox.showinfo("加载成功", f"已加载{self.template_type.upper()}模板：\n{os.path.basename(path)}")

    def load_roster(self):
        """加载人员台账（新增列名更新）"""
        path = filedialog.askopenfilename(
            filetypes=[("Excel文件", "*.xlsx")])
        if path:
            try:
                wb = load_workbook(path)
                ws = wb.active
                self.roster_data = [[cell.value for cell in row] for row in ws.rows]
                
                # 更新列名并刷新下拉框
                self.column_names = [str(cell) for cell in self.roster_data[0]] if self.roster_data else []
                self._update_comboboxes()
                
                messagebox.showinfo("加载成功", f"成功加载{len(self.roster_data)-1}条人员数据")
            except Exception as e:
                messagebox.showerror("错误", f"读取Excel失败：{str(e)}")

    def _update_comboboxes(self):
        """更新所有下拉框选项"""
        for _, combobox in self.replace_rules:
            combobox['values'] = self.column_names
            if combobox.get() not in self.column_names:
                combobox.set('')  # 清除非法的当前值

    def select_output_dir(self):
        """手动选择输出目录"""
        path = filedialog.askdirectory()
        if path:
            self.output_dir = path
            messagebox.showinfo("路径设置", f"输出目录已设置为：\n{path}")

    # ------------------- 核心处理函数 -------------------
    def process_template(self, rules, output_dir, row_data):
        """统一模板处理方法"""
        if self.template_type == "docx":
            return self.process_word_template(rules, output_dir, row_data)
        elif self.template_type == "xlsx":
            return self.process_excel_template(rules, output_dir, row_data)
        return False

    def process_word_template(self, rules, output_dir, row_data):
        """处理Word模板"""
        try:
            doc = Document(self.template_path)
            date_str = datetime.now().strftime("%Y%m%d")
            
            # 处理段落和表格
            for para in doc.paragraphs:
                self.replace_runs(para.runs, rules, row_data, date_str)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.replace_runs(para.runs, rules, row_data, date_str)
            
            filename = self.generate_filename(row_data, date_str)
            doc.save(os.path.join(output_dir, filename))
            return True
        except Exception as e:
            messagebox.showerror("错误", f"Word生成失败：{str(e)}")
            return False

    def process_excel_template(self, rules, output_dir, row_data):
        """新增Excel模板处理"""
        try:
            wb = load_workbook(self.template_path)
            ws = wb.active
            date_str = datetime.now().strftime("%Y%m%d")

            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        text = cell.value
                        for src, target in rules.items():
                            value = row_data.get(target, "")
                            if value is None:
                                value = ""
                            text = text.replace(src, str(value))
                        cell.value = text.replace("{日期}", date_str)
            
            filename = self.generate_filename(row_data, date_str)
            wb.save(os.path.join(output_dir, filename.replace(".docx", ".xlsx")))
            return True
        except Exception as e:
            messagebox.showerror("错误", f"Excel生成失败：{str(e)}")
            return False

    def replace_runs(self, runs, rules, row_data, date_str):
        """替换文本内容"""
        for run in runs:
            text = run.text
            for src, target in rules.items():
                value = row_data.get(target, "")
                if value is None:
                    value = ""
                text = text.replace(src, str(value))
            run.text = text.replace("{日期}", date_str)

    def generate_filename(self, row_data, date_str):
        """生成文件名"""
        filename = self.filename_pattern.get()
        filename = filename.replace("{日期}", date_str)
        for key in row_data:
            if f"{{{key}}}" in filename:
                filename = filename.replace(f"{{{key}}}", str(row_data[key]))
        return f"{filename}.docx" if self.template_type == "docx" else f"{filename}.xlsx"

    # ------------------- 线程控制 -------------------
    def start_generation(self):
        """启动生成线程"""
        if self.validate_inputs():
            threading.Thread(target=self.run_generation).start()

    def run_generation(self):
        """执行生成任务"""
        try:
            # 修改获取规则的方式
            rules = {
                src.get(): tgt.get() 
                for src, tgt in self.replace_rules 
                if src.get() and tgt.get() in self.column_names  # 增加有效性验证
            }
            
            output_path = self.prepare_output_dir()
            total = len(self.roster_data) - 1
            
            for i, row in enumerate(self.roster_data[1:], 1):
                row_data = dict(zip(self.roster_data[0], row))
                self.process_template(rules, output_path, row_data)
                self.progress['value'] = (i / total) * 100
                self.update()
            
            messagebox.showinfo("完成", f"成功生成{total}份档案！")
            if messagebox.askyesno("打开目录", "是否打开输出目录？"):
                os.startfile(output_path)
            self.progress['value'] = 0
        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{str(e)}")
            self.progress['value'] = 0

    def prepare_output_dir(self):
        """准备输出目录"""
        if self.auto_dir.get():
            dir_name = f"一人一档_{datetime.now():%Y%m%d}"
            path = os.path.join(self.output_dir, dir_name)
            os.makedirs(path, exist_ok=True)
            return path
        return self.output_dir

    def validate_inputs(self):
        """验证输入完整性"""
        errors = []
        if not self.template_path:
            errors.append("请选择模板文件")
        if len(self.roster_data) < 2:
            errors.append("台账数据不足")
        if not self.output_dir:
            errors.append("请设置输出目录")
        
        if errors:
            messagebox.showerror("输入错误", "\n".join(errors))
            return False
        return True

if __name__ == "__main__":
    BatchGenerator()
